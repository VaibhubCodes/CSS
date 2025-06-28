# import docx
# import io
# from django.http import JsonResponse 
# from django.shortcuts import get_object_or_404
# from .models import UserFile,OCRResult,FileCategory
# from storage_management.models import UserStorage
# import boto3
# from django.conf import settings

# def extract_text_from_document(file_obj, file_extension):
#     """Extract text from various document formats"""
#     try:
#         if file_extension == 'txt':
#             # For text files
#             content = file_obj.read()
#             try:
#                 # Try UTF-8 first
#                 return content.decode('utf-8')
#             except UnicodeDecodeError:
#                 # Fall back to latin-1 if UTF-8 fails
#                 return content.decode('latin-1')
                
#         elif file_extension == 'docx':
#             # For Word documents
#             doc = docx.Document(io.BytesIO(file_obj.read()))
#             return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
#         return None
        
#     except Exception as e:
#         raise Exception(f"Error extracting text from {file_extension} file: {str(e)}")

# # # Update process_document_ocr view
# # def process_document_ocr(request, file_id):
# #     try:
# #         user_file = get_object_or_404(UserFile, id=file_id)
# #         file_name = user_file.file.name
# #         file_extension = file_name.split('.')[-1].lower()
        
# #         # Set default category as Personal
# #         default_category, _ = FileCategory.objects.get_or_create(
# #             name='Personal',
# #             defaults={'is_default': True}
# #         )
# #         user_file.category = default_category
# #         user_file.save()

# #         def categorize_file(text_content):
# #             """Helper function to categorize the file based on extracted text."""
# #             categorization_service = FileCategorizationService()
# #             analysis = categorization_service.analyze_file_content(text_content)
            
# #             if analysis['confidence'] >= 40:
# #                 category, _ = FileCategory.objects.get_or_create(
# #                     name=analysis['category'],
# #                     defaults={'is_default': True}
# #                 )
# #                 user_file.category = category
# #                 user_file.save()
# #                 return category.name, analysis
# #             return 'Personal', analysis

# #         # Handle text-based files (txt, docx, md)
# #         if file_extension in ['txt', 'docx', 'md']:
# #             try:
# #                 file_content = extract_text_from_document(user_file.file, file_extension)
# #                 if file_content:
# #                     category_name, analysis = categorize_file(file_content)
# #                     ocr_result, created = OCRResult.objects.update_or_create(
# #                         file=user_file,
# #                         defaults={
# #                             'status': 'completed',
# #                             'text_content': file_content
# #                         }
# #                     )
# #                     return JsonResponse({
# #                         'status': 'completed',
# #                         'text': file_content.split('\n'),
# #                         'category': category_name,
# #                         'analysis': analysis
# #                     })
# #             except Exception as e:
# #                 return JsonResponse({
# #                     'error': f'Error processing file: {str(e)}',
# #                     'category': 'Personal'
# #                 }, status=500)

# #         # Initialize AWS Textract client
# #         textract_client = boto3.client(
# #             'textract',
# #             aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
# #             aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY,
# #             region_name=settings.AWS_S3_REGION_NAME
# #         )

# #         document_path = user_file.file.name

# #         try:
# #             # Handle images (jpg, jpeg, png)
# #             if file_extension in ['jpg', 'jpeg', 'png']:
# #                 response = textract_client.detect_document_text(
# #                     Document={
# #                         'S3Object': {
# #                             'Bucket': settings.AWS_STORAGE_BUCKET_NAME,
# #                             'Name': document_path
# #                         }
# #                     }
# #                 )
                
# #                 extracted_text = [item['Text'] for item in response['Blocks'] if item['BlockType'] == 'LINE']
# #                 text_content = '\n'.join(extracted_text)
                
# #                 # Categorize content
# #                 category_name, analysis = categorize_file(text_content)
                
# #                 ocr_result, created = OCRResult.objects.update_or_create(
# #                     file=user_file,
# #                     defaults={
# #                         'status': 'completed',
# #                         'text_content': text_content
# #                     }
# #                 )
                
# #                 return JsonResponse({
# #                     'status': 'completed',
# #                     'text': extracted_text,
# #                     'category': category_name,
# #                     'analysis': analysis
# #                 })

# #             # Handle PDFs
# #             elif file_extension == 'pdf':
# #                 response = textract_client.start_document_analysis(
# #                     DocumentLocation={
# #                         'S3Object': {
# #                             'Bucket': settings.AWS_STORAGE_BUCKET_NAME,
# #                             'Name': document_path
# #                         }
# #                     },
# #                     FeatureTypes=['TABLES', 'FORMS']
# #                 )
                
# #                 job_id = response['JobId']
                
# #                 # Create or update OCR result
# #                 ocr_result, created = OCRResult.objects.update_or_create(
# #                     file=user_file,
# #                     defaults={
# #                         'status': 'processing',
# #                         'job_id': job_id
# #                     }
# #                 )
                
# #                 return JsonResponse({
# #                     'status': 'processing',
# #                     'job_id': job_id
# #                 })
            
# #             else:
# #                 return JsonResponse({
# #                     'error': 'Unsupported file type',
# #                     'category': 'Personal'
# #                 }, status=400)

# #         except textract_client.exceptions.InvalidS3ObjectException:
# #             return JsonResponse({
# #                 'error': 'File not accessible in S3',
# #                 'category': 'Personal'
# #             }, status=400)
# #         except textract_client.exceptions.UnsupportedDocumentException:
# #             return JsonResponse({
# #                 'error': 'Document format not supported',
# #                 'category': 'Personal'
# #             }, status=400)
# #         except Exception as e:
# #             return JsonResponse({
# #                 'error': str(e),
# #                 'category': 'Personal'
# #             }, status=500)

# #     except Exception as e:
# #         return JsonResponse({
# #             'error': str(e),
# #             'category': 'Personal'
# #         }, status=500)


# class FileCategorizationService:
#     CATEGORY_KEYWORDS = {
#         'Notes': [
#             'note', 'notes', 'memo', 'reminder', 'todo', 'task',
#             'meeting notes', 'summary', 'key points', 'highlights',
#             'agenda', 'minutes', 'checklist', 'observations',
#             'review notes', 'lecture notes'
#         ],
#         'Personal': [
#             'personal', 'family', 'home', 'private', 'diary',
#             'vacation', 'travel', 'birthday', 'anniversary',
#             'journal', 'memories', 'personal goals', 'bucket list',
#             'life events', 'hobbies', 'recipes', 'fitness'
#         ],
#         'Banking': [
#             'bank', 'account', 'transaction', 'credit', 'debit',
#             'statement', 'balance', 'loan', 'emi', 'payment',
#             'transfer', 'deposit', 'withdrawal', 'ifsc', 'bank branch',
#             'interest', 'savings', 'cheque', 'passbook', 'net banking',
#             'upi', 'credit card', 'debit card', 'banking'
#         ],
#         'Professional': [
#             'resume', 'cv', 'work', 'job', 'office', 'business',
#             'client', 'project', 'proposal', 'contract', 'agreement',
#             'meeting', 'professional', 'corporate', 'company',
#             'employment', 'salary', 'appraisal', 'performance review',
#             'offer letter', 'resignation', 'portfolio', 'work experience',
#             'interview', 'background'
#         ],
#         'Education': [
#             'school', 'college', 'university', 'course', 'study',
#             'assignment', 'exam', 'grade', 'result', 'certificate',
#             'degree', 'diploma', 'student', 'professor', 'lecture',
#             'syllabus', 'curriculum', 'academic', 'education',
#             'semester', 'research', 'thesis', 'dissertation',
#             'classroom', 'training', 'workshop', 'seminar','syllabus'
#         ],
#         'Weblinks': [
#             'http', 'https', 'www', '.com', '.org', '.edu',
#             'website', 'url', 'link', 'web', 'site', 'portal',
#             'online', 'internet', 'webpage', 'browser', 'bookmark',
#             'domain', 'web address', 'hyperlink', 'html'
#         ],
#         'Medical': [
#             'hospital', 'doctor', 'prescription', 'medical', 'health',
#             'medicine', 'patient', 'clinic', 'report', 'test',
#             'diagnosis', 'treatment', 'pharmacy', 'appointment',
#             'healthcare', 'insurance', 'pathology', 'symptoms',
#             'medication', 'consultation', 'laboratory', 'dental',
#             'x-ray', 'vaccination', 'immunization'
#         ],
#         'Visiting_Cards': [
#             'tel:', 'phone:', 'email:', '@', 'contact', 'address:',
#             'mobile', 'website:', 'designation', 'company name',
#             'business card', 'visiting card', 'contact details',
#             'phone number', 'fax:', 'cell:', 'office address',
#             'linkedin', 'social media', 'professional profile'
#         ],
#         'Investments': [
#             'invest', 'stock', 'share', 'mutual fund', 'dividend',
#             'portfolio', 'returns', 'investment', 'equity', 'bond',
#             'demat', 'trading', 'nse', 'bse', 'market', 'profit',
#             'securities', 'holdings', 'financial', 'assets',
#             'commodities', 'forex', 'cryptocurrency', 'bitcoin',
#             'gold', 'real estate', 'property', 'capital gains'
#         ]
#     }

#     def get_category(self, text_content, default_category='Personal'):
#         """
#         Analyze text content and return the most likely category using weighted scoring
#         """
#         if not text_content:
#             return default_category

#         text_lower = text_content.lower()
        
#         # Score categories with weights
#         category_scores = {}
        
#         for category, keywords in self.CATEGORY_KEYWORDS.items():
#             score = 0
#             for keyword in keywords:
#                 # Count occurrences
#                 count = text_lower.count(keyword.lower())
                
#                 # Add weight based on keyword specificity
#                 if len(keyword) > 8:  # Longer keywords are more specific
#                     score += count * 2
#                 else:
#                     score += count
                
#                 # Add extra weight for exact matches
#                 if f" {keyword.lower()} " in f" {text_lower} ":
#                     score += 3
                    
#                 # Add extra weight for keywords in title/filename
#                 if keyword.lower() in text_lower.split('\n')[0].lower():
#                     score += 5

#             if score > 0:
#                 category_scores[category] = score

#         # Calculate confidence scores
#         if category_scores:
#             max_score = max(category_scores.values())
#             if max_score >= 3:  # Minimum threshold for categorization
#                 return max(category_scores.items(), key=lambda x: x[1])[0]
            
#         return default_category

#     def analyze_file_content(self, text_content):
#         """
#         Analyze file content and return detailed categorization info
#         """
#         if not text_content:
#             return {'category': 'Personal', 'confidence': 0, 'matches': []}

#         text_lower = text_content.lower()
#         category_matches = {}
        
#         for category, keywords in self.CATEGORY_KEYWORDS.items():
#             matches = []
#             for keyword in keywords:
#                 count = text_lower.count(keyword.lower())
#                 if count > 0:
#                     matches.append({
#                         'keyword': keyword,
#                         'count': count,
#                         'exact_match': f" {keyword.lower()} " in f" {text_lower} "
#                     })
#             if matches:
#                 category_matches[category] = matches

#         # Get primary category
#         category = self.get_category(text_content)
        
#         # Calculate confidence
#         confidence = 0
#         if category in category_matches:
#             total_matches = sum(match['count'] for match in category_matches[category])
#             exact_matches = sum(1 for match in category_matches[category] if match['exact_match'])
#             confidence = min((total_matches + exact_matches * 2) / 10, 1) * 100

#         return {
#             'category': category,
#             'confidence': confidence,
#             'matches': category_matches
#         }

# def create_default_categories():
#     """Create default categories if they don't exist"""
#     default_categories = [
#         ('Personal', 'Personal documents and information'),
#         ('Banking', 'Banking statements and financial documents'),
#         ('Professional', 'Work-related documents'),
#         ('Education', 'Academic certificates and educational documents'),
#         ('Medical', 'Healthcare-related documents'),
#         ('Visiting Cards', 'Business cards and contact information'),
#         ('Investments', 'Investment certificates and documents'),
#         ('Miscellaneous', 'Uncategorized documents')
#     ]
    
#     for name, description in default_categories:
#         FileCategory.objects.get_or_create(
#             name=name,
#             defaults={
#                 'description': description,
#                 'is_default': True
#             }
#         )


import docx
import io
import boto3
from django.conf import settings
from django.shortcuts import get_object_or_404
from .models import UserFile, OCRResult, FileCategory
import requests

def download_file_from_s3(s3_key):
    """Download file content from S3"""
    try:
        s3_client = boto3.client(
            's3',
            aws_access_key_id=settings.AWS_ACCESS_KEY_ID,
            aws_secret_access_key=settings.AWS_SECRET_ACCESS_KEY,
            region_name=settings.AWS_S3_REGION_NAME
        )
        
        response = s3_client.get_object(
            Bucket=settings.AWS_STORAGE_BUCKET_NAME,
            Key=s3_key
        )
        return response['Body'].read()
    except Exception as e:
        print(f"Error downloading file from S3: {str(e)}")
        raise

def extract_text_from_document(s3_key, file_extension):
    """Extract text from various document formats stored in S3"""
    try:
        # Download file content from S3
        file_content = download_file_from_s3(s3_key)
        
        if file_extension == 'txt':
            # For text files
            try:
                return file_content.decode('utf-8')
            except UnicodeDecodeError:
                return file_content.decode('latin-1')
                
        elif file_extension == 'docx':
            # For Word documents
            doc = docx.Document(io.BytesIO(file_content))
            return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
        elif file_extension == 'md':
            # For Markdown files
            try:
                return file_content.decode('utf-8')
            except UnicodeDecodeError:
                return file_content.decode('latin-1')
            
        return None
        
    except Exception as e:
        raise Exception(f"Error extracting text from {file_extension} file: {str(e)}")

class FileCategorizationService:
    CATEGORY_KEYWORDS = {
        'Professional': [
            # Core professional terms
            'resume', 'cv', 'curriculum vitae', 'work', 'job', 'office', 'business',
            'cover letter', 'application letter', 'employment', 'career',
            
            # Job application terms
            'dear hiring manager', 'dear sir', 'dear madam', 'application',
            'opportunity', 'candidate', 'hire', 'hiring', 'recruit', 'position',
            'role', 'vacancy', 'opening', 'interview', 'qualified', 'qualifications',
            
            # Work experience terms
            'experience', 'skills', 'achievements', 'responsibilities', 'duties',
            'accomplishments', 'projects', 'managed', 'developed', 'led', 'created',
            'implemented', 'improved', 'increased', 'reduced', 'streamlined',
            
            # Professional relationships
            'client', 'customer', 'colleague', 'team', 'supervisor', 'manager',
            'director', 'executive', 'professional', 'corporate', 'company',
            'organization', 'department', 'division',
            
            # Business terms
            'project', 'proposal', 'contract', 'agreement', 'meeting',
            'presentation', 'report', 'analysis', 'strategy', 'planning',
            'budget', 'revenue', 'profit', 'sales', 'marketing',
            
            # Employment terms
            'salary', 'compensation', 'benefits', 'promotion', 'appraisal',
            'performance review', 'offer letter', 'resignation', 'termination',
            'notice period', 'reference', 'recommendation',
            
            # Closing terms
            'sincerely', 'regards', 'best regards', 'yours faithfully',
            'thank you for your consideration', 'look forward to hearing',
            'available for interview', 'discuss further'
        ],
        'Education': [
            'school', 'college', 'university', 'course', 'study', 'education',
            'assignment', 'exam', 'grade', 'result', 'certificate', 'transcript',
            'degree', 'diploma', 'bachelor', 'master', 'phd', 'doctorate',
            'student', 'professor', 'teacher', 'instructor', 'lecture', 'class',
            'syllabus', 'curriculum', 'academic', 'semester', 'term', 'year',
            'research', 'thesis', 'dissertation', 'paper', 'publication',
            'classroom', 'training', 'workshop', 'seminar', 'conference',
            'scholarship', 'financial aid', 'tuition', 'enrollment', 'admission',
            'gpa', 'honors', 'dean list', 'graduation', 'commencement'
        ],
        'Banking': [
            'bank', 'banking', 'account', 'checking', 'savings', 'transaction',
            'credit', 'debit', 'statement', 'balance', 'loan', 'mortgage',
            'emi', 'payment', 'transfer', 'deposit', 'withdrawal', 'atm',
            'ifsc', 'routing number', 'bank branch', 'interest', 'apr',
            'cheque', 'check', 'passbook', 'net banking', 'online banking',
            'upi', 'credit card', 'debit card', 'mastercard', 'visa',
            'overdraft', 'minimum balance', 'service charge', 'fee',
            'investment', 'mutual fund', 'fixed deposit', 'recurring deposit',
            'swift code', 'wire transfer', 'direct deposit', 'ach'
        ],
        'Medical': [
            'hospital', 'clinic', 'doctor', 'physician', 'nurse', 'medical',
            'health', 'healthcare', 'medicine', 'medication', 'prescription',
            'patient', 'diagnosis', 'treatment', 'therapy', 'surgery',
            'appointment', 'consultation', 'examination', 'checkup',
            'insurance', 'copay', 'deductible', 'claim', 'coverage',
            'pathology', 'laboratory', 'test', 'results', 'blood work',
            'x-ray', 'mri', 'ct scan', 'ultrasound', 'biopsy',
            'symptoms', 'condition', 'disease', 'illness', 'injury',
            'vaccination', 'immunization', 'shot', 'vaccine',
            'dental', 'dentist', 'orthodontist', 'optometrist', 'cardiologist'
        ],
        'Personal': [
            'personal', 'family', 'home', 'private', 'diary', 'journal',
            'vacation', 'holiday', 'travel', 'trip', 'birthday', 'anniversary',
            'wedding', 'celebration', 'party', 'event', 'memories',
            'personal goals', 'bucket list', 'life events', 'milestone',
            'hobbies', 'interests', 'recreation', 'leisure', 'entertainment',
            'recipes', 'cooking', 'fitness', 'exercise', 'workout', 'gym',
            'relationships', 'friendship', 'dating', 'marriage', 'children',
            'pets', 'house', 'apartment', 'moving', 'utilities', 'bills'
        ],
        'Legal': [
            'legal', 'law', 'attorney', 'lawyer', 'court', 'judge', 'lawsuit',
            'contract', 'agreement', 'terms', 'conditions', 'clause',
            'settlement', 'litigation', 'trial', 'hearing', 'deposition',
            'affidavit', 'notary', 'witness', 'testimony', 'evidence',
            'plaintiff', 'defendant', 'jurisdiction', 'statute', 'regulation',
            'compliance', 'violation', 'penalty', 'fine', 'damages',
            'intellectual property', 'copyright', 'trademark', 'patent',
            'will', 'estate', 'trust', 'probate', 'inheritance'
        ],
        'Notes': [
            'note', 'notes', 'memo', 'reminder', 'todo', 'task', 'list',
            'meeting notes', 'summary', 'key points', 'highlights', 'takeaways',
            'agenda', 'minutes', 'checklist', 'observations', 'thoughts',
            'review notes', 'lecture notes', 'study notes', 'research notes',
            'ideas', 'brainstorm', 'draft', 'outline', 'plan', 'strategy'
        ],
        'Investments': [
            'invest', 'investment', 'portfolio', 'stock', 'stocks', 'share',
            'shares', 'equity', 'bond', 'bonds', 'mutual fund', 'etf',
            'dividend', 'returns', 'yield', 'capital gains', 'profit', 'loss',
            'demat', 'trading', 'broker', 'brokerage', 'commission',
            'nse', 'bse', 'nasdaq', 'nyse', 'market', 'bull', 'bear',
            'securities', 'holdings', 'asset', 'assets', 'allocation',
            'diversification', 'risk', 'volatility', 'performance',
            'commodities', 'gold', 'silver', 'oil', 'forex', 'currency',
            'cryptocurrency', 'bitcoin', 'ethereum', 'blockchain',
            'real estate', 'property', 'reit', 'retirement', '401k', 'ira'
        ]
    }

    def get_category(self, text_content, default_category='Miscellaneous'):
        """Analyze text content and return the most likely category"""
        if not text_content or len(text_content.strip()) < 10:
            return default_category

        text_lower = text_content.lower()
        category_scores = {}
        
        for category, keywords in self.CATEGORY_KEYWORDS.items():
            score = 0
            keyword_matches = 0
            
            for keyword in keywords:
                # Count occurrences
                count = text_lower.count(keyword.lower())
                if count > 0:
                    keyword_matches += 1
                    
                    # Add weight based on keyword specificity and length
                    if len(keyword) > 10:  # Very specific keywords
                        score += count * 4
                    elif len(keyword) > 6:  # Moderately specific
                        score += count * 3
                    else:
                        score += count * 2
                    
                    # Bonus for exact word matches
                    if f" {keyword.lower()} " in f" {text_lower} ":
                        score += 5
                        
                    # Extra bonus for keywords in first part of document
                    first_part = text_lower[:500]
                    if keyword.lower() in first_part:
                        score += 3

            # Apply multiplier based on number of different keywords matched
            if keyword_matches > 0:
                diversity_multiplier = min(1 + (keyword_matches * 0.1), 2.0)
                category_scores[category] = score * diversity_multiplier

        if category_scores:
            max_score = max(category_scores.values())
            # Lower threshold for better categorization
            if max_score >= 5:  # Reduced from 10
                return max(category_scores.items(), key=lambda x: x[1])[0]
            
        return default_category

    def analyze_file_content(self, text_content):
        """Analyze file content and return detailed categorization info"""
        if not text_content or len(text_content.strip()) < 10:
            return {'category': 'Miscellaneous', 'confidence': 0, 'matches': []}

        text_lower = text_content.lower()
        category_matches = {}
        category_scores = {}
        
        for category, keywords in self.CATEGORY_KEYWORDS.items():
            matches = []
            total_score = 0
            
            for keyword in keywords:
                count = text_lower.count(keyword.lower())
                if count > 0:
                    keyword_score = count * (len(keyword) // 3 + 1)  # Length-based scoring
                    if f" {keyword.lower()} " in f" {text_lower} ":
                        keyword_score += 5  # Exact match bonus
                    
                    total_score += keyword_score
                    matches.append({
                        'keyword': keyword,
                        'count': count,
                        'exact_match': f" {keyword.lower()} " in f" {text_lower} ",
                        'score': keyword_score
                    })
            
            if matches:
                category_matches[category] = matches
                category_scores[category] = total_score

        # Get primary category
        category = self.get_category(text_content)
        
        # Calculate confidence
        confidence = 0
        if category in category_scores and category_scores:
            max_score = max(category_scores.values())
            total_score = sum(category_scores.values())
            
            if total_score > 0:
                # Base confidence on relative score
                relative_score = category_scores[category] / max_score
                confidence = min(relative_score * 100, 100)

        return {
            'category': category,
            'confidence': confidence,
            'matches': category_matches,
            'scores': category_scores
        }
    
def create_default_categories():
    """Create default categories if they don't exist"""
    default_categories = [
        ('Personal', 'Personal documents and information'),
        ('Banking', 'Banking statements and financial documents'),
        ('Professional', 'Work-related documents'),
        ('Education', 'Academic certificates and educational documents'),
        ('Medical', 'Healthcare-related documents'),
        ('Visiting Cards', 'Business cards and contact information'),
        ('Investments', 'Investment certificates and documents'),
        ('Miscellaneous', 'Uncategorized documents')
    ]
    
    for name, description in default_categories:
        FileCategory.objects.get_or_create(
            name=name,
            defaults={
                'description': description,
                'is_default': True
            }
        )


